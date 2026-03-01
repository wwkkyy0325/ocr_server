# -*- coding: utf-8 -*-

try:
    from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QTableWidget, QTableWidgetItem, 
                               QPushButton, QHBoxLayout, QHeaderView, QFileDialog, QAbstractItemView, QShortcut)
    from PyQt5.QtCore import Qt, pyqtSignal
    from PyQt5.QtGui import QColor, QBrush, QFont, QKeySequence, QGuiApplication
    from app.ui.dialogs.glass_dialogs import GlassMessageDialog
    PYQT_AVAILABLE = True
except ImportError:
    PYQT_AVAILABLE = False

import csv
import os

class ResultTableWidget(QWidget):
    request_preferred_width = pyqtSignal(int)
    """
    Widget to display OCR results in a structured table format.
    Supports merged cells (rowspan/colspan) and export to XLSX/CSV.
    """
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.layout = QVBoxLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)
        
        # Toolbar
        self.toolbar_layout = QHBoxLayout()
        self.btn_export_xlsx = QPushButton("导出 Excel")
        self.btn_export_csv = QPushButton("导出 CSV")
        self.btn_export_word = QPushButton("导出 Word")
        self.btn_auto_fit = QPushButton("恢复布局")
        
        self.toolbar_layout.addWidget(self.btn_export_xlsx)
        self.toolbar_layout.addWidget(self.btn_export_csv)
        self.toolbar_layout.addWidget(self.btn_export_word)
        self.toolbar_layout.addWidget(self.btn_auto_fit)
        self.toolbar_layout.addStretch()
        
        self.layout.addLayout(self.toolbar_layout)
        
        # Table Widget
        self.table = QTableWidget()
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.table.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Interactive)
        self.table.horizontalHeader().sectionDoubleClicked.connect(self._on_horizontal_header_double_clicked)
        self.table.verticalHeader().sectionDoubleClicked.connect(self._on_vertical_header_double_clicked)
        self.table.setAlternatingRowColors(True)
        # 监听选中变化，用于清理残留的悬停样式
        self.table.itemSelectionChanged.connect(self._on_selection_changed)
        
        # Add table widget to layout
        self.layout.addWidget(self.table)

        # Mapping from block index to list of (row, col)
        self._block_to_cells = {}
        self._hovered_block_index = -1
        
        # Connections
        self.btn_export_xlsx.clicked.connect(self._export_to_xlsx)
        self.btn_export_csv.clicked.connect(self._export_to_csv)
        self.btn_export_word.clicked.connect(self._export_to_word)
        self.btn_auto_fit.clicked.connect(self._on_auto_fit_clicked)
        self.copy_shortcut = QShortcut(QKeySequence.Copy, self.table)
        self.copy_shortcut.activated.connect(self.copy_selection)
        
        # Styling
        self.btn_export_xlsx.setStyleSheet("")
        self.btn_export_csv.setStyleSheet("")
        self.btn_export_word.setStyleSheet("")
        self.btn_auto_fit.setStyleSheet("")
        # 基础文字颜色 + 更明显的选中样式（悬停由逻辑控制）
        self.table.setStyleSheet("""
            QTableWidget {
                color: #E0E0E0;
            }
            QTableWidget::item:selected {
                background-color: rgba(24, 144, 255, 190);
                color: #FFFFFF;
            }
        """)
        
        self._initial_column_widths = []
        self._initial_row_heights = []
        self._initial_total_width = 0
        self._export_basename = ""
        
    def set_block_mapping(self, block_blocks):
        """
        可选：从 TextBlockListWidget 同步 block -> cells 映射
        目前内部自建映射即可，此接口保留拓展。
        """
        self._block_to_cells = block_blocks or {}

    def set_export_basename(self, basename: str):
        """
        设置导出文件名的基础名（不含扩展名），通常使用图片文件名
        """
        self._export_basename = basename or ""
        
    def set_hovered_block(self, block_index):
        """
        从 ImageViewer 悬停事件同步块高亮：高亮整块区域，不残留旧块。
        如果 block_index 为 -1，则执行 clear_hover()。
        """
        # 如果传入的 block_index 为 -1，说明鼠标移出了任何块，清理所有悬停
        if block_index == -1:
            self.clear_hover()
            return
            
        # 如果新块和旧块相同，无需操作
        if self._hovered_block_index == block_index:
            return
            
        # 1. 强制清理旧 hover（无论之前状态如何，先恢复原状）
        self.clear_hover()
            
        # 2. 设置新 hover
        self._hovered_block_index = block_index
        if self._block_to_cells and block_index != -1:
            cells = self._block_to_cells.get(block_index, [])
            for r, c in cells:
                item = self.table.item(r, c)
                if item:
                    # 如果单元格已经被选中，则保持选中样式，不覆盖为悬停样式
                    if item.isSelected():
                        continue
                        
                    # 悬停状态：比选中略浅一点的主题蓝色背景 + 白字
                    item.setBackground(QBrush(QColor(35, 80, 150, 200)))
                    item.setForeground(QBrush(QColor(255, 255, 255)))
        
        # 强制刷新视口以应用更改
        self.table.viewport().update()

    def clear_hover(self):
        """清除所有悬停高亮，恢复单元格默认样式（除非被选中）"""
        # 注意：这里我们不仅清理 _hovered_block_index 对应的，还要遍历整个表格确保没有残留
        # 为了性能，我们还是先只清理记录的 _hovered_block_index
        # 如果用户反馈还有残留，可能需要遍历所有 cells（虽然性能较差但稳妥）
        
        if self._hovered_block_index != -1 and self._block_to_cells:
            cells = self._block_to_cells.get(self._hovered_block_index, [])
            for r, c in cells:
                item = self.table.item(r, c)
                if item:
                    # 如果单元格被选中，则保留选中状态，不恢复默认
                    if item.isSelected():
                        continue
                        
                    # 恢复默认样式
                    item.setBackground(QBrush(Qt.NoBrush))
                    item.setForeground(QBrush(QColor(224, 224, 224)))
        
        # 额外加一道保险：如果出现逻辑错误导致 _hovered_block_index 丢失但界面残留
        # 可以在必要时遍历清理（暂不开启，视性能而定）
        
        self._hovered_block_index = -1
        self.table.viewport().update()

    def _on_selection_changed(self):
        """
        当表格选中项发生变化时触发
        需要清理那些“曾经被选中但现在不再选中”的项的背景色（如果它们还残留着悬停色）
        """
        # 遍历所有单元格可能太慢，我们利用 Qt 的机制
        # Qt 的 selectionModel changed 会自动处理选中状态的绘制
        # 但我们需要确保那些“失去选中”的单元格，如果没有被悬停，就恢复透明背景
        
        # 简单暴力的做法：强制刷新整个表格的非选中项背景
        # 但为了性能，我们只处理当前可见区域或利用 clear_hover
        
        # 实际上，当选中发生变化时，如果之前的选中项变成了非选中，它可能会因为之前的悬停逻辑残留背景
        # 我们可以在这里强制清理一次 hover，因为点击操作通常意味着 hover 目标的改变或结束
        self.clear_hover()
        
        # 另外，对于所有非选中的 item，强制恢复背景为 NoBrush
        # 这可能比较耗时，但能彻底解决残留
        # 优化：只遍历已知的 block cells？不，用户可能框选了多个区域
        
        # 让我们尝试只在鼠标操作导致的选中变化时清理
        # 只要发生选中变化，就说明交互状态改变了，清理悬停是合理的
        pass
        
    def set_data(self, ocr_results):
        """
        Populate table from OCR results.
        ocr_results: List of dicts, expected to contain 'table_info'
        """
        self.table.clear()
        self.table.setRowCount(0)
        self.table.setColumnCount(0)
        self._initial_column_widths = []
        self._initial_row_heights = []
        self._initial_total_width = 0
        self._block_to_cells = {}
        
        if not ocr_results:
            return

        # 1. Determine dimensions
        max_row = 0
        max_col = 0
        
        cells = []
        
        for item in ocr_results:
            table_info = item.get('table_info')
            if not table_info:
                continue
                
            # Ensure values are integers and handle None
            try:
                row = int(table_info.get('row') if table_info.get('row') is not None else 0)
                col = int(table_info.get('col') if table_info.get('col') is not None else 0)
                rowspan = int(table_info.get('rowspan') if table_info.get('rowspan') is not None else 1)
                colspan = int(table_info.get('colspan') if table_info.get('colspan') is not None else 1)
            except (ValueError, TypeError):
                # Fallback if conversion fails
                row, col = 0, 0
                rowspan, colspan = 1, 1
            
            max_row = max(max_row, row + rowspan - 1)
            max_col = max(max_col, col + colspan - 1)
            
            cells.append({
                'text': item.get('text', ''),
                'row': row,
                'col': col,
                'rowspan': rowspan,
                'colspan': colspan,
                'is_header': table_info.get('is_header', False)
            })
            
        if not cells:
            return
            
        # 1.1 Remove empty rows and columns
        # Identify non-empty rows and columns
        non_empty_rows = set()
        non_empty_cols = set()
        
        block_map = {}
        for cell in cells:
            if cell['text'].strip():
                # Mark range of rows/cols as non-empty
                for r in range(cell['row'], cell['row'] + cell['rowspan']):
                    non_empty_rows.add(r)
                for c in range(cell['col'], cell['col'] + cell['colspan']):
                    non_empty_cols.add(c)
        
        # If all empty, maybe show nothing or keep as is? 
        # Let's keep at least one cell if everything is empty to avoid confusion, 
        # or just proceed which will result in 0x0 table if strictly empty.
        # But usually we want to see the structure even if empty?
        # User request: "去除空行空列" (remove empty rows/cols).
        # So if a row has absolutely no text in any cell, remove it.
        
        # Create mapping from old indices to new indices
        sorted_rows = sorted(list(non_empty_rows))
        sorted_cols = sorted(list(non_empty_cols))
        
        row_map = {old_idx: new_idx for new_idx, old_idx in enumerate(sorted_rows)}
        col_map = {old_idx: new_idx for new_idx, old_idx in enumerate(sorted_cols)}
        
        # Filter and remap cells
        new_cells = []
        new_max_row = len(sorted_rows) - 1
        new_max_col = len(sorted_cols) - 1
        
        if not sorted_rows or not sorted_cols:
             # Case: All cells are empty text. 
             # We should probably check if we want to show structure or nothing.
             # If completely empty, return empty table.
             return

        for cell in cells:
            # Check if cell is within non-empty rows/cols
            # A merged cell might span across empty and non-empty rows/cols.
            # We need to adjust its start and span.
            
            # Simplified approach: Only keep cells that start in a non-empty row/col?
            # Or better: check intersection.
            
            # Logic:
            # New Row = row_map[original_row] (if original_row is in non_empty)
            # But what if original_row is empty but part of a rowspan?
            # Actually, if a row is "empty", it means NO cell has text in that row.
            # So if a cell has text, its rows are in non_empty_rows.
            
            # However, a cell might span multiple rows. If it has text, ALL its spanned rows 
            # might not be necessary? No, if it has text, usually we consider the cell "occupies" those rows.
            # But my logic above `for r in range... non_empty_rows.add(r)` marks ALL spanned rows as non-empty.
            # So if a cell has text, all rows it spans are preserved.
            
            # What if a cell has NO text?
            # It might still need to be preserved if it's in a row that has OTHER text?
            # Yes. We only remove a row if *entire* row has no text.
            # My logic: `if cell['text'].strip(): ... add(r)`
            # This logic correctly identifies rows that contain at least one non-empty cell.
            # BUT: What about a row that has only empty cells? It will not be in `non_empty_rows`.
            # So we drop it.
            
            # Now we need to remap cells.
            # If a cell (empty or not) belongs to a preserved row/col, we keep it.
            # Be careful with spans.
            # If a cell is empty, but sits in a preserved row, we keep it.
            
            # Wait, `non_empty_rows` currently only contains rows that have *text-containing cells*.
            # Rows with only empty cells are excluded.
            # This is correct for "remove empty rows".
            
            # Issue: What if an empty cell spans across a preserved row and a removed row?
            # e.g. Row 0: "A", Row 1: (empty). Cell at (0,0) spans 2 rows.
            # Row 0 is preserved. Row 1 is removed.
            # Cell (0,0) new span should be 1.
            
            # Algorithm to remap:
            # Calculate new row index, col index, rowspan, colspan.
            
            r_start = cell['row']
            r_end = cell['row'] + cell['rowspan'] # exclusive
            c_start = cell['col']
            c_end = cell['col'] + cell['colspan'] # exclusive
            
            # Find which of these rows/cols survive
            surviving_rows = [r for r in range(r_start, r_end) if r in row_map]
            surviving_cols = [c for c in range(c_start, c_end) if c in col_map]
            
            if not surviving_rows or not surviving_cols:
                continue
                
            new_r = row_map[surviving_rows[0]]
            new_c = col_map[surviving_cols[0]]
            new_rs = len(surviving_rows)
            new_cs = len(surviving_cols)
            
            # Update cell
            cell['row'] = new_r
            cell['col'] = new_c
            cell['rowspan'] = new_rs
            cell['colspan'] = new_cs
            new_cells.append(cell)

        cells = new_cells
        max_row = new_max_row
        max_col = new_max_col
            
        self.table.setRowCount(max_row + 1)
        self.table.setColumnCount(max_col + 1)
        
        # 2. Populate
        # Note: QTableWidget needs items for all cells, but for merged cells we set spans.
        # We must be careful not to overwrite spans.
        
        # Reset spans
        self.table.clearSpans()
        
        for cell in cells:
            r, c = cell['row'], cell['col']
            rs, cs = cell['rowspan'], cell['colspan']
            text = cell['text']
            is_header = cell['is_header']
            
            item = QTableWidgetItem(text)
            
            # Styling
            if is_header:
                item.setBackground(QColor(220, 235, 255))
                font = item.font()
                font.setBold(True)
                item.setFont(font)
            
            self.table.setItem(r, c, item)
            # 使用原 table_info 行列做 block key，保证同一逻辑单元格的多行文本作为一块
            key = (cell['row'], cell['col'])
            lst = block_map.get(key)
            if lst is None:
                block_map[key] = [(r, c)]
            else:
                lst.append((r, c))
            
            if rs > 1 or cs > 1:
                self.table.setSpan(r, c, rs, cs)

        self._block_to_cells = block_map
        self._auto_fit_all(resize_splitter=True, remember=True)

    def _auto_fit_all(self, resize_splitter, remember=False):
        fm = self.table.fontMetrics()
        col_count = self.table.columnCount()
        row_count = self.table.rowCount()
        total_width = 0
        for col in range(col_count):
            max_width = 0
            header_item = self.table.horizontalHeaderItem(col)
            if header_item:
                rect = fm.boundingRect(header_item.text())
                max_width = rect.width() + 16
            for row in range(row_count):
                item = self.table.item(row, col)
                if not item:
                    continue
                text = item.text()
                if not text:
                    continue
                lines = text.splitlines() or [""]
                line_width = 0
                for line in lines:
                    rect = fm.boundingRect(line)
                    if rect.width() > line_width:
                        line_width = rect.width()
                width = line_width + 16
                if width > max_width:
                    max_width = width
            if max_width <= 0:
                max_width = self.table.horizontalHeader().sectionSizeHint(col)
            self.table.setColumnWidth(col, max_width)
            total_width += max_width

        for row in range(row_count):
            max_height = 0
            for col in range(col_count):
                item = self.table.item(row, col)
                if not item:
                    continue
                text = item.text()
                if not text:
                    continue
                lines = text.splitlines() or [""]
                height = fm.lineSpacing() * len(lines) + 8
                if height > max_height:
                    max_height = height
            if max_height <= 0:
                max_height = self.table.verticalHeader().sectionSizeHint(row)
            self.table.setRowHeight(row, max_height)

        if remember:
            self._initial_column_widths = [self.table.columnWidth(c) for c in range(col_count)]
            self._initial_row_heights = [self.table.rowHeight(r) for r in range(row_count)]
            frame = self.table.frameWidth() * 2
            self._initial_total_width = total_width + frame + self.table.verticalHeader().width()

        if resize_splitter and total_width > 0:
            total = self._initial_total_width if remember and self._initial_total_width > 0 else total_width + self.table.frameWidth() * 2 + self.table.verticalHeader().width()
            self.request_preferred_width.emit(total)

    def _on_horizontal_header_double_clicked(self, section):
        if section < 0:
            return
        fm = self.table.fontMetrics()
        max_width = 0
        header_item = self.table.horizontalHeaderItem(section)
        if header_item:
            rect = fm.boundingRect(header_item.text())
            max_width = rect.width() + 16
        for row in range(self.table.rowCount()):
            item = self.table.item(row, section)
            if not item:
                continue
            text = item.text()
            if not text:
                continue
            lines = text.splitlines() or [""]
            line_width = 0
            for line in lines:
                rect = fm.boundingRect(line)
                if rect.width() > line_width:
                    line_width = rect.width()
            w = line_width + 16
            if w > max_width:
                max_width = w
        if max_width > 0:
            self.table.setColumnWidth(section, max_width)
            total_width = 0
            for col in range(self.table.columnCount()):
                total_width += self.table.columnWidth(col)
            if total_width > 0:
                frame = self.table.frameWidth() * 2
                total = total_width + frame + self.table.verticalHeader().width()
                self.request_preferred_width.emit(total)

    def _on_vertical_header_double_clicked(self, section):
        if section < 0:
            return
        header = self.table.verticalHeader()
        header.setSectionResizeMode(section, QHeaderView.Interactive)
        fm = self.table.fontMetrics()
        max_height = 0
        for col in range(self.table.columnCount()):
            item = self.table.item(section, col)
            if not item:
                continue
            text = item.text()
            if not text:
                continue
            lines = text.splitlines() or [""]
            h = fm.lineSpacing() * len(lines) + 8
            if h > max_height:
                max_height = h
        if max_height > 0:
            self.table.setRowHeight(section, max_height)
            
    def _on_auto_fit_clicked(self):
        if not self._initial_column_widths and not self._initial_row_heights:
            self._auto_fit_all(resize_splitter=True, remember=True)
            return
        col_count = self.table.columnCount()
        row_count = self.table.rowCount()
        for c in range(min(col_count, len(self._initial_column_widths))):
            self.table.setColumnWidth(c, self._initial_column_widths[c])
        for r in range(min(row_count, len(self._initial_row_heights))):
            self.table.setRowHeight(r, self._initial_row_heights[r])
        if self._initial_total_width > 0:
            self.request_preferred_width.emit(self._initial_total_width)
                
    def _export_to_xlsx(self):
        """Export current table to XLSX"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            dlg = GlassMessageDialog(
                self,
                title="错误",
                text="未安装 openpyxl 库，无法导出 Excel。",
                buttons=[("ok", "确定")],
            )
            dlg.exec_()
            return

        default_name = f"{self._export_basename}.xlsx" if self._export_basename else ""
        path, _ = QFileDialog.getSaveFileName(self, "导出 Excel", default_name, "Excel Files (*.xlsx)")
        if not path:
            return
            
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            
            # Populate
            rows = self.table.rowCount()
            cols = self.table.columnCount()
            
            # We need to handle spans in Excel too
            # And also, QTableWidget might have empty items where span covers.
            
            # Track processed cells to avoid overwriting spanned areas
            # Actually openpyxl handles merge_cells, we just need to write to top-left
            
            for r in range(rows):
                for c in range(cols):
                    item = self.table.item(r, c)
                    if item:
                        cell = ws.cell(row=r+1, column=c+1, value=item.text())
                        
                        # Style
                        if item.background().color() == QColor(220, 235, 255):
                            cell.fill = PatternFill(start_color="DCEBFF", end_color="DCEBFF", fill_type="solid")
                            cell.font = Font(bold=True)
                        
                        cell.alignment = Alignment(wrap_text=True, vertical='center')
            
            # Apply Spans
            # QTableWidget spans are (row, col, rowspan, colspan)
            # Openpyxl needs start_row, start_col, end_row, end_col
            
            # We iterate cells again or store spans?
            # It's better to iterate our internal data if possible, but reading from table is WYSIWYG
            
            # There isn't a direct "get all spans" in QTableWidget API easily exposed without loop
            # But we can check span for each cell.
            
            # Optimization: check top-left of spans
            for r in range(rows):
                for c in range(cols):
                    rs = self.table.rowSpan(r, c)
                    cs = self.table.columnSpan(r, c)
                    if rs > 1 or cs > 1:
                        ws.merge_cells(start_row=r+1, start_column=c+1, end_row=r+rs, end_column=c+cs)
            
            wb.save(path)
            dlg_ok = GlassMessageDialog(
                self,
                title="成功",
                text=f"表格已导出至: {path}",
                buttons=[("ok", "确定")],
            )
            dlg_ok.exec_()
            
        except Exception as e:
            dlg_err = GlassMessageDialog(
                self,
                title="导出失败",
                text=f"导出过程中发生错误: {e}",
                buttons=[("ok", "确定")],
            )
            dlg_err.exec_()
            
    def _export_to_csv(self):
        """Export to CSV (ignores styling and merges - flattens or duplicates?)"""
        default_name = f"{self._export_basename}.csv" if self._export_basename else ""
        path, _ = QFileDialog.getSaveFileName(self, "导出 CSV", default_name, "CSV Files (*.csv)")
        if not path:
            return
            
        try:
            with open(path, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                rows = self.table.rowCount()
                cols = self.table.columnCount()
                
                for r in range(rows):
                    row_data = []
                    for c in range(cols):
                        item = self.table.item(r, c)
                        text = item.text() if item else ""
                        row_data.append(text)
                    writer.writerow(row_data)
                    
            dlg_ok2 = GlassMessageDialog(
                self,
                title="成功",
                text=f"表格已导出至: {path}",
                buttons=[("ok", "确定")],
            )
            dlg_ok2.exec_()
        except Exception as e:
            dlg_err2 = GlassMessageDialog(
                self,
                title="导出失败",
                text=f"导出过程中发生错误: {e}",
                buttons=[("ok", "确定")],
            )
            dlg_err2.exec_()

    def _export_to_word(self):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            dlg = GlassMessageDialog(
                self,
                title="错误",
                text="未安装 python-docx 库，无法导出 Word。",
                buttons=[("ok", "确定")],
            )
            dlg.exec_()
            return

        default_name = f"{self._export_basename}.docx" if self._export_basename else ""
        path, _ = QFileDialog.getSaveFileName(self, "导出 Word", default_name, "Word Documents (*.docx)")
        if not path:
            return

        def set_cell_shading(cell, fill_hex):
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            tc_pr.append(shd)

        def style_cell_text(cell, bold=False, align=None):
            for p in cell.paragraphs:
                if align is not None:
                    p.alignment = align
                for run in p.runs:
                    run.font.bold = bold
                    run.font.size = Pt(10.5)
                    run.font.name = "宋体"
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.get_or_add_rFonts()
                    rfonts.set(qn("w:eastAsia"), "宋体")

        try:
            doc = Document()
            # 设置中文字体，确保显示正常
            try:
                doc.styles['Normal'].font.name = u'宋体'
                doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            except Exception:
                pass
            
            rows = self.table.rowCount()
            cols = self.table.columnCount()

            for r in range(rows):
                row_texts = []
                # 标记该行是否全是空文本
                has_content = False
                
                for c in range(cols):
                    # 检查单元格跨度，如果当前单元格被上方单元格覆盖（rowSpan），则视为空
                    # 但 QTableWidget 的 item 在被覆盖位置通常为空，所以直接取 text 即可
                    item = self.table.item(r, c)
                    text = item.text() if item else ""
                    
                    if text.strip():
                        has_content = True
                    
                    # 简单地将每个单元格视为一个文本块
                    # 如果是空单元格，保留为空字符串，以便 join 时产生间隔
                    row_texts.append(text)

                # 如果整行都没有内容，是否跳过？
                # 用户说“换行保持好”，空行可能是有意义的段落间隔，保留。
                # 但全是空串的 join 结果是空白字符串，add_paragraph 会加一个空行。
                
                # 使用制表符或空格连接
                # 这里使用4个空格模拟间距，或者使用制表符
                # 用户提到“中间空格”，我们用空格
                line_str = "    ".join(row_texts).rstrip()
                
                if line_str:
                    doc.add_paragraph(line_str)
                elif has_content: 
                    # 理论上进不来，但为了保险
                    pass
                else:
                    # 空行，添加空段落
                    doc.add_paragraph("")

            doc.save(path)
            dlg_ok = GlassMessageDialog(
                self,
                title="成功",
                text=f"Word 已导出至: {path}",
                buttons=[("ok", "确定")],
            )
            dlg_ok.exec_()
        except Exception as e:
            dlg_err = GlassMessageDialog(
                self,
                title="导出失败",
                text=f"导出过程中发生错误: {e}",
                buttons=[("ok", "确定")],
            )
            dlg_err.exec_()

    def copy_selection(self):
        ranges = self.table.selectedRanges()
        if not ranges:
            return

        lines = []
        for table_range in ranges:
            top = table_range.topRow()
            bottom = table_range.bottomRow()
            left = table_range.leftColumn()
            right = table_range.rightColumn()

            for r in range(top, bottom + 1):
                row_texts = []
                for c in range(left, right + 1):
                    item = self.table.item(r, c)
                    text = item.text() if item else ""
                    row_texts.append(text)
                lines.append("\t".join(row_texts))

            lines.append("")

        if lines and lines[-1] == "":
            lines.pop()

        text = "\n".join(lines)
        QGuiApplication.clipboard().setText(text)
